from __future__ import annotations

import io
import math
import re
import textwrap
from pathlib import Path
from typing import Any, Iterable

import ezdxf
import fitz
from ezdxf.addons import Importer
from ezdxf import bbox
from ezdxf.addons.drawing import Frontend, RenderContext, layout, pymupdf
from ezdxf.addons.drawing.config import (
    BackgroundPolicy,
    ColorPolicy,
    Configuration,
)
from ezdxf.math import BoundingBox2d, Vec2
from pypdf import PdfReader, PdfWriter


SKIP_RENDER_ENTITY_TYPES = {"IMAGE", "OLE2FRAME", "PDFUNDERLAY", "DGNUNDERLAY", "DWFUNDERLAY"}

STAMP_SIGNATURE_TARGET_Y = {
    "Разработал": -1.0,
    "Проверил": -6.2,
    "Н.контроль": -15.5,
    "Хаустов": -1.0,
}
STAMP_SURINOV_TOP_Y_THRESHOLD = -12.0
STAMP_SURINOV_TOP_TARGET_Y = -6.2
STAMP_SURINOV_BOTTOM_TARGET_Y = -15.5


def _pdf_render_configuration(*, faithful: bool = False) -> Configuration:
    if faithful:
        # Сохраняем цвета и оформление чертежа как в DWG.
        return Configuration(
            color_policy=ColorPolicy.COLOR,
            background_policy=BackgroundPolicy.WHITE,
        )
    return Configuration(
        color_policy=ColorPolicy.BLACK,
        background_policy=BackgroundPolicy.WHITE,
    )


def _pdf_layout_settings(*, faithful: bool = False) -> layout.Settings:
    return layout.Settings(
        fit_page=True,
        page_alignment=layout.PageAlignment.MIDDLE_CENTER,
        crop_at_margins=False,
        min_stroke_width=0.05,
        fixed_stroke_width=0.12,
    )


def render_dxf_to_pdf(dxf_path: Path, output_pdf_path: Path, *, faithful: bool = False) -> Path:
    if faithful:
        return render_dxf_to_pdf_faithful(dxf_path, output_pdf_path)

    dxf_path = Path(dxf_path)
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    document = ezdxf.readfile(dxf_path)
    skip_image_refs = dxf_path.name.lower() != "check_result.dxf"
    if skip_image_refs:
        _remove_skipped_render_entities(document)
    writer = PdfWriter()
    rendered_pages = 0
    modelspace = document.modelspace()
    sheet_boxes = _modelspace_sheet_boxes(modelspace)
    adjust_stamp_text = dxf_path.name.lower() != "check_result.dxf"

    if sheet_boxes:
        output_pdf_path.write_bytes(
            _render_modelspace_sheet_images(
                document,
                modelspace,
                sheet_boxes,
                adjust_stamp_text=adjust_stamp_text,
            )
        )
        return output_pdf_path
    else:
        rendered_layout_names: set[str] = set()
        for dxf_layout in _renderable_layouts(document):
            layout_sheet_boxes = _layout_sheet_boxes(dxf_layout)
            if layout_sheet_boxes:
                for sheet_box in layout_sheet_boxes:
                    try:
                        pdf_bytes = _render_layout(document, dxf_layout, render_box=sheet_box)
                    except ValueError:
                        continue
                    reader = PdfReader(io.BytesIO(pdf_bytes))
                    for page in reader.pages:
                        writer.add_page(page)
                        rendered_pages += 1
                rendered_layout_names.add(dxf_layout.name)
                continue

            try:
                pdf_bytes = _render_layout(document, dxf_layout)
            except ValueError:
                continue
            pdf_bytes = _split_wide_pdf_pages(pdf_bytes)
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                writer.add_page(page)
                rendered_pages += 1
                rendered_layout_names.add(dxf_layout.name)

        if rendered_pages == 0 and "Model" not in rendered_layout_names:
            try:
                pdf_bytes = _render_layout(document, modelspace)
                pdf_bytes = _split_wide_pdf_pages(pdf_bytes)
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)
                    rendered_pages += 1
            except ValueError:
                pass

    if rendered_pages == 0:
        raise RuntimeError(f"Не удалось найти листы для PDF-экспорта в {dxf_path.name}.")

    with output_pdf_path.open("wb") as stream:
        writer.write(stream)
    return output_pdf_path


def render_dxf_to_pdf_faithful(dxf_path: Path, output_pdf_path: Path) -> Path:
    """Экспорт DXF в PDF без правок содержимого чертежа."""
    dxf_path = Path(dxf_path)
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    document = ezdxf.readfile(dxf_path)
    writer = PdfWriter()
    rendered_pages = 0
    modelspace = document.modelspace()
    sheet_boxes = _modelspace_sheet_boxes(modelspace)

    if sheet_boxes:
        output_pdf_path.write_bytes(
            _render_modelspace_sheet_images(
                document,
                modelspace,
                sheet_boxes,
                faithful=True,
            )
        )
        return output_pdf_path

    for dxf_layout in _renderable_layouts(document):
        layout_sheet_boxes = _layout_sheet_boxes(dxf_layout)
        if layout_sheet_boxes:
            for sheet_box in layout_sheet_boxes:
                try:
                    pdf_bytes = _render_layout(
                        document,
                        dxf_layout,
                        render_box=sheet_box,
                        faithful=True,
                    )
                except ValueError:
                    continue
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)
                    rendered_pages += 1
            continue

        try:
            pdf_bytes = _render_layout(document, dxf_layout, faithful=True)
        except ValueError:
            continue
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
            rendered_pages += 1

    if rendered_pages == 0:
        try:
            pdf_bytes = _render_layout(document, modelspace, faithful=True)
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                writer.add_page(page)
                rendered_pages += 1
        except ValueError:
            pass

    if rendered_pages == 0:
        raise RuntimeError(f"Не удалось найти листы для PDF-экспорта в {dxf_path.name}.")

    with output_pdf_path.open("wb") as stream:
        writer.write(stream)
    return output_pdf_path


def merge_pdfs(pdf_paths: list[Path], output_pdf_path: Path) -> Path:
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()

    for path in pdf_paths:
        if not Path(path).exists():
            continue
        reader = PdfReader(str(path))
        for page in reader.pages:
            writer.add_page(page)

    if not writer.pages:
        raise RuntimeError("Нет PDF-файлов для объединения.")

    with output_pdf_path.open("wb") as stream:
        writer.write(stream)
    return output_pdf_path


def render_tu_docx_to_pdf(docx_path: Path, output_pdf_path: Path) -> Path:
    from docx import Document

    docx_path = Path(docx_path)
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(docx_path)
    lines = _docx_lines(document)
    if not lines:
        raise RuntimeError("DOCX-ТУ не содержит текста для вставки в PDF.")

    pdf = fitz.open()
    font_path = _find_text_font()
    font = fitz.Font(fontfile=str(font_path)) if font_path else fitz.Font("tiro")
    font_name = "TuText"
    font_size = 10.0
    line_height = 13.4
    margin_x = 42.0
    margin_y = 42.0
    page_width = 595.0
    page_height = 842.0
    max_width = page_width - margin_x * 2
    max_y = page_height - margin_y

    page = pdf.new_page(width=page_width, height=page_height)
    y = margin_y
    for raw_line in lines:
        wrapped = _wrap_pdf_text(raw_line, font, font_size, max_width)
        if not wrapped:
            y += line_height
            continue
        for line in wrapped:
            if y > max_y:
                page = pdf.new_page(width=page_width, height=page_height)
                y = margin_y
            page.insert_text(
                (margin_x, y),
                line,
                fontsize=font_size,
                fontname=font_name,
                fontfile=str(font_path) if font_path else None,
                color=(0, 0, 0),
            )
            y += line_height
        y += 2.0

    pdf.save(output_pdf_path)
    pdf.close()
    return output_pdf_path


def merge_project_pdfs(
    note_pdf_path: Path,
    plan_pdf_path: Path | None,
    tu_pdf_path: Path | None,
    output_pdf_path: Path,
) -> Path:
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    note_reader = PdfReader(str(note_pdf_path))
    writer = PdfWriter()

    section_2_index = _find_pdf_page_index(
        note_reader,
        ("Раздел 2", "Проект полосы отвода"),
        fallback_index=9,
    )
    section_3_index = _find_pdf_page_index(
        note_reader,
        ("Раздел 3", "Технологические"),
        fallback_index=13,
    )
    tu_reader = _optional_pdf_reader(tu_pdf_path)
    plan_reader = _optional_pdf_reader(plan_pdf_path)
    note_page_count = len(note_reader.pages)
    tu_insert_index = section_2_index if section_2_index < note_page_count else min(1, note_page_count)
    plan_insert_index = section_3_index if section_3_index < note_page_count else note_page_count
    tu_inserted = False
    plan_inserted = False

    for index, page in enumerate(note_reader.pages):
        if tu_reader is not None and index == tu_insert_index:
            _add_reader_pages(writer, tu_reader)
            tu_inserted = True
        if plan_reader is not None and index == plan_insert_index:
            _add_reader_pages(writer, plan_reader)
            plan_inserted = True
        writer.add_page(page)

    if tu_reader is not None and not tu_inserted:
        _add_reader_pages(writer, tu_reader)
    if plan_reader is not None and not plan_inserted:
        _add_reader_pages(writer, plan_reader)

    if not writer.pages:
        raise RuntimeError("Нет PDF-файлов для объединения.")

    with output_pdf_path.open("wb") as stream:
        writer.write(stream)
    return output_pdf_path


def _docx_lines(document: Any) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = _normalize_pdf_text(paragraph.text)
        if text:
            lines.append(text)
        elif lines and lines[-1] != "":
            lines.append("")

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_pdf_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                lines.append(row_text)
        lines.append("")
    return lines


def _normalize_pdf_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def _wrap_pdf_text(text: str, font: fitz.Font, font_size: float, max_width: float) -> list[str]:
    if not text:
        return []
    lines: list[str] = []
    current = ""
    for word in text.split():
        candidate = f"{current} {word}".strip()
        if font.text_length(candidate, fontsize=font_size) <= max_width:
            current = candidate
            continue
        if current:
            lines.append(current)
        if font.text_length(word, fontsize=font_size) <= max_width:
            current = word
            continue
        chunks = textwrap.wrap(word, width=42, break_long_words=True, break_on_hyphens=False)
        lines.extend(chunks[:-1])
        current = chunks[-1] if chunks else ""
    if current:
        lines.append(current)
    return lines


def _find_text_font() -> Path | None:
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/Library/Fonts/Times New Roman.ttf"),
        Path("/Library/Fonts/Arial.ttf"),
        Path("C:/Windows/Fonts/times.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _renderable_layouts(document: ezdxf.EzDxf) -> list[ezdxf.layouts.Layout]:
    paper_layouts = [
        dxf_layout
        for dxf_layout in document.layouts
        if dxf_layout.name.lower() != "model" and _has_finite_extents(dxf_layout)
    ]
    if paper_layouts:
        return paper_layouts

    modelspace = document.modelspace()
    return [modelspace] if _has_finite_extents(modelspace) else []


def _remove_skipped_render_entities(document: ezdxf.EzDxf) -> None:
    for dxf_layout in document.layouts:
        for entity in list(dxf_layout):
            if entity.dxftype() in SKIP_RENDER_ENTITY_TYPES or _is_oversized_paperspace_viewport(entity, dxf_layout):
                try:
                    dxf_layout.delete_entity(entity)
                except Exception:
                    continue


def _is_oversized_paperspace_viewport(entity: ezdxf.entities.DXFEntity, dxf_layout: ezdxf.layouts.Layout) -> bool:
    if entity.dxftype() != "VIEWPORT" or dxf_layout.name.lower() == "model":
        return False
    try:
        extents = bbox.extents([entity])
    except Exception:
        return False
    if not extents.has_data:
        return False
    width = abs(float(extents.size.x))
    height = abs(float(extents.size.y))
    return width > 600 or height > 430


def _render_layout(
    document: ezdxf.EzDxf,
    dxf_layout: ezdxf.layouts.Layout,
    render_box: BoundingBox2d | None = None,
    *,
    faithful: bool = False,
) -> bytes:
    page = _page_for_layout(dxf_layout, render_box=render_box)
    settings = _pdf_layout_settings(faithful=faithful)
    config = _pdf_render_configuration(faithful=faithful)
    backend_class = pymupdf.PyMuPdfBackend
    backend = backend_class()
    Frontend(RenderContext(document), backend, config=config).draw_layout(dxf_layout)
    return backend.get_pdf_bytes(page, settings=settings, render_box=render_box)


def _render_modelspace_sheet_images(
    document: ezdxf.EzDxf,
    modelspace: ezdxf.layouts.Modelspace,
    sheet_boxes: list[BoundingBox2d],
    adjust_stamp_text: bool = False,
    *,
    faithful: bool = False,
) -> bytes:
    writer = PdfWriter()
    entity_boxes = _entity_boxes(modelspace)

    for sheet_box in sheet_boxes:
        sheet_entities = _entities_for_sheet(modelspace, entity_boxes, sheet_box)
        if not sheet_entities:
            continue

        sheet_document = ezdxf.new(document.dxfversion)
        sheet_document.units = document.units
        importer = Importer(document, sheet_document)
        importer.import_entities(_expand_render_entities(sheet_entities), sheet_document.modelspace())
        importer.finalize()

        if faithful:
            pdf_bytes = _render_layout(
                sheet_document,
                sheet_document.modelspace(),
                render_box=sheet_box,
                faithful=True,
            )
        else:
            if adjust_stamp_text:
                _adjust_stamp_signature_text(sheet_document)
            _add_multileader_text_overlays(sheet_entities, sheet_document.modelspace())
            pdf_bytes = _render_layout(
                sheet_document,
                sheet_document.modelspace(),
                render_box=sheet_box,
            )

        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)

    if not writer.pages:
        raise RuntimeError("Не удалось отрисовать листы modelspace в PDF.")

    stream = io.BytesIO()
    writer.write(stream)
    return stream.getvalue()


def _expand_render_entities(entities: list[ezdxf.entities.DXFEntity]) -> list[ezdxf.entities.DXFEntity]:
    expanded: list[ezdxf.entities.DXFEntity] = []
    for entity in entities:
        if entity.dxftype() in SKIP_RENDER_ENTITY_TYPES:
            continue
        if entity.dxftype() == "MULTILEADER":
            try:
                virtual_entities = list(entity.virtual_entities())
                expanded.extend(
                    virtual
                    for virtual in virtual_entities
                    if virtual.dxftype() not in {"TEXT", "MTEXT", "HATCH"}
                )
                continue
            except Exception:
                pass
        if entity.dxftype() == "ACAD_TABLE":
            try:
                expanded.extend(entity.virtual_entities())
                continue
            except Exception:
                pass
        expanded.append(entity)
    return expanded


def _expand_render_entities_for_pdf(
    entities: list[ezdxf.entities.DXFEntity],
) -> list[ezdxf.entities.DXFEntity]:
    expanded: list[ezdxf.entities.DXFEntity] = []
    table_text_by_position: dict[tuple[float, float], ezdxf.entities.DXFEntity] = {}

    for entity in entities:
        if entity.dxftype() in SKIP_RENDER_ENTITY_TYPES:
            continue
        if entity.dxftype() == "MULTILEADER":
            try:
                virtual_entities = list(entity.virtual_entities())
                expanded.extend(
                    virtual
                    for virtual in virtual_entities
                    if virtual.dxftype() not in {"TEXT", "MTEXT", "HATCH"}
                )
                continue
            except Exception:
                pass
        if entity.dxftype() == "ACAD_TABLE":
            try:
                for virtual in _dedupe_table_virtual_entities(entity.virtual_entities()):
                    position_key = _text_entity_position_key(virtual)
                    if position_key is None:
                        expanded.append(virtual)
                        continue
                    _keep_longer_text_at_position(table_text_by_position, position_key, virtual)
                continue
            except Exception:
                pass
        expanded.append(entity)

    expanded.extend(table_text_by_position.values())
    return expanded


def _keep_longer_text_at_position(
    text_by_position: dict[tuple[float, float], ezdxf.entities.DXFEntity],
    position_key: tuple[float, float],
    entity: ezdxf.entities.DXFEntity,
) -> None:
    existing = text_by_position.get(position_key)
    if existing is None or _text_entity_plain_length(entity) > _text_entity_plain_length(existing):
        text_by_position[position_key] = entity


def _dedupe_table_virtual_entities(
    virtual_entities: Iterable[ezdxf.entities.DXFEntity],
) -> list[ezdxf.entities.DXFEntity]:
    """Оставляет один текст на ячейку таблицы — самый полный вариант MTEXT."""
    deduped: list[ezdxf.entities.DXFEntity] = []
    text_by_position: dict[tuple[float, float], list[ezdxf.entities.DXFEntity]] = {}

    for virtual in virtual_entities:
        position_key = _text_entity_position_key(virtual)
        if position_key is None:
            deduped.append(virtual)
            continue
        text_by_position.setdefault(position_key, []).append(virtual)

    for group in text_by_position.values():
        if len(group) == 1:
            deduped.append(group[0])
            continue
        deduped.append(max(group, key=_text_entity_plain_length))
    return deduped


def _text_entity_position_key(
    entity: ezdxf.entities.DXFEntity,
) -> tuple[float, float] | None:
    if entity.dxftype() == "TEXT":
        insert = getattr(entity.dxf, "insert", None)
    elif entity.dxftype() == "MTEXT":
        insert = getattr(entity.dxf, "insert", None)
    else:
        return None
    if insert is None:
        return None
    return (round(float(insert.x), 2), round(float(insert.y), 2))


def _text_entity_plain_length(entity: ezdxf.entities.DXFEntity) -> int:
    if entity.dxftype() == "TEXT":
        text = str(getattr(entity.dxf, "text", "") or "")
    elif entity.dxftype() == "MTEXT":
        text = str(getattr(entity, "text", "") or "")
    else:
        return 0
    return len(_plain_text(text))


def _adjust_stamp_signature_text(document: ezdxf.EzDxf) -> None:
    for block in document.blocks:
        for entity in block:
            if entity.dxftype() not in {"TEXT", "MTEXT"}:
                continue
            text = getattr(entity.dxf, "text", "") if entity.dxftype() == "TEXT" else getattr(entity, "text", "")
            insert = getattr(entity.dxf, "insert", None)
            if insert is None:
                continue
            target_y = _stamp_signature_target_y(_plain_text(text), insert.y)
            if target_y is None:
                continue
            entity.dxf.insert = (insert.x, target_y, insert.z)


def _stamp_signature_target_y(text: str, insert_y: float | None = None) -> float | None:
    if text == "Суринов" and insert_y is not None:
        if insert_y <= STAMP_SURINOV_TOP_Y_THRESHOLD:
            return STAMP_SURINOV_BOTTOM_TARGET_Y
        return STAMP_SURINOV_TOP_TARGET_Y
    return STAMP_SIGNATURE_TARGET_Y.get(text)


def _plain_text(text: str) -> str:
    result = re.sub(r"\\[A-Za-z][^;{}\\]*(?:;)?", " ", text)
    result = result.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", result).strip()


def _add_multileader_text_overlays(
    entities: list[ezdxf.entities.DXFEntity],
    target_layout: ezdxf.layouts.Modelspace,
) -> None:
    for entity in entities:
        if entity.dxftype() != "MULTILEADER":
            continue
        try:
            virtual_entities = list(entity.virtual_entities())
        except Exception:
            continue
        for virtual in virtual_entities:
            if virtual.dxftype() == "TEXT":
                text = getattr(virtual.dxf, "text", "")
                if not text:
                    continue
                attribs = {
                    "insert": virtual.dxf.insert,
                    "height": getattr(virtual.dxf, "height", 2.5),
                    "rotation": getattr(virtual.dxf, "rotation", 0.0),
                    "layer": getattr(virtual.dxf, "layer", "0"),
                    "color": 7,
                }
                style = getattr(virtual.dxf, "style", None)
                if style:
                    attribs["style"] = style
                copied = target_layout.add_text(text, dxfattribs=attribs)
                copied.dxf.halign = getattr(virtual.dxf, "halign", 0)
                copied.dxf.valign = getattr(virtual.dxf, "valign", 0)
            elif virtual.dxftype() == "MTEXT":
                text = getattr(virtual, "text", "")
                if not text:
                    continue
                attribs = {
                    "insert": virtual.dxf.insert,
                    "char_height": getattr(virtual.dxf, "char_height", 2.5),
                    "rotation": getattr(virtual.dxf, "rotation", 0.0),
                    "layer": getattr(virtual.dxf, "layer", "0"),
                    "color": 7,
                }
                style = getattr(virtual.dxf, "style", None)
                if style:
                    attribs["style"] = style
                target_layout.add_mtext(text, dxfattribs=attribs)


def _optional_pdf_reader(path: Path | None) -> PdfReader | None:
    if path is None:
        return None
    path = Path(path)
    if path.suffix.lower() != ".pdf" or not path.exists():
        return None
    return PdfReader(str(path))


def _add_reader_pages(writer: PdfWriter, reader: PdfReader) -> None:
    for page in reader.pages:
        writer.add_page(page)


def _find_pdf_page_index(reader: PdfReader, markers: tuple[str, ...], fallback_index: int) -> int:
    normalized_markers = tuple(_normalize_pdf_text(marker) for marker in markers)
    for index, page in enumerate(reader.pages):
        try:
            text = _normalize_pdf_text(page.extract_text() or "")
        except Exception:
            continue
        if all(marker in text for marker in normalized_markers):
            return index
    if 0 <= fallback_index < len(reader.pages):
        return fallback_index
    return len(reader.pages)


def _normalize_pdf_text(text: str) -> str:
    return " ".join(text.casefold().replace("\xa0", " ").split())


def _split_wide_pdf_pages(pdf_bytes: bytes) -> bytes:
    document = fitz.open(stream=pdf_bytes, filetype="pdf")
    if len(document) != 1:
        return pdf_bytes

    page = document[0]
    pixmap = page.get_pixmap(dpi=100, alpha=False)
    width = pixmap.width
    height = pixmap.height
    split_x = _find_vertical_split(pixmap)
    if split_x is None:
        return pdf_bytes

    left_bbox = _ink_bbox(pixmap, 0, split_x)
    right_bbox = _ink_bbox(pixmap, split_x, width)
    if left_bbox is None or right_bbox is None:
        return pdf_bytes
    if not _looks_like_two_sheets(left_bbox, right_bbox, width, height):
        return pdf_bytes

    output = fitz.open()
    for ink_bbox in (left_bbox, right_bbox):
        clip = _pixel_bbox_to_page_rect(ink_bbox, page.rect, width, height)
        clip = _expanded_rect(clip, page.rect, margin=6)
        target_page = _standard_page_for_rect(clip)
        output_page = output.new_page(width=target_page.width, height=target_page.height)
        output_page.show_pdf_page(output_page.rect, document, 0, clip=clip)

    return output.tobytes(deflate=True, garbage=4)


def _find_vertical_split(pixmap: fitz.Pixmap) -> int | None:
    width = pixmap.width
    height = pixmap.height
    samples = pixmap.samples
    stride = pixmap.n
    threshold = max(5, int(height * 0.002))
    column_counts: list[int] = []

    for x in range(width):
        count = 0
        for y in range(height):
            index = (y * width + x) * stride
            if _is_ink(samples[index], samples[index + 1], samples[index + 2]):
                count += 1
        column_counts.append(count)

    search_start = int(width * 0.15)
    search_end = int(width * 0.85)
    best_run: tuple[int, int] | None = None
    run_start: int | None = None

    for x in range(search_start, search_end):
        if column_counts[x] <= threshold:
            if run_start is None:
                run_start = x
        elif run_start is not None:
            if best_run is None or x - run_start > best_run[1] - best_run[0]:
                best_run = (run_start, x)
            run_start = None

    if run_start is not None and (best_run is None or search_end - run_start > best_run[1] - best_run[0]):
        best_run = (run_start, search_end)

    if best_run is None or best_run[1] - best_run[0] < 5:
        return None
    return (best_run[0] + best_run[1]) // 2


def _ink_bbox(pixmap: fitz.Pixmap, x_start: int, x_end: int) -> tuple[int, int, int, int] | None:
    width = pixmap.width
    height = pixmap.height
    samples = pixmap.samples
    stride = pixmap.n
    min_x = width
    min_y = height
    max_x = -1
    max_y = -1

    for x in range(max(0, x_start), min(width, x_end)):
        for y in range(height):
            index = (y * width + x) * stride
            if _is_ink(samples[index], samples[index + 1], samples[index + 2]):
                min_x = min(min_x, x)
                min_y = min(min_y, y)
                max_x = max(max_x, x)
                max_y = max(max_y, y)

    if max_x < min_x or max_y < min_y:
        return None
    return (min_x, min_y, max_x, max_y)


def _is_ink(red: int, green: int, blue: int) -> bool:
    return red < 245 or green < 245 or blue < 245


def _looks_like_two_sheets(
    left_bbox: tuple[int, int, int, int],
    right_bbox: tuple[int, int, int, int],
    page_width: int,
    page_height: int,
) -> bool:
    left_width = left_bbox[2] - left_bbox[0] + 1
    right_width = right_bbox[2] - right_bbox[0] + 1
    left_height = left_bbox[3] - left_bbox[1] + 1
    right_height = right_bbox[3] - right_bbox[1] + 1
    if min(left_width, right_width) < page_width * 0.15:
        return False
    if min(left_height, right_height) < page_height * 0.35:
        return False
    return True


def _pixel_bbox_to_page_rect(
    bbox_pixels: tuple[int, int, int, int],
    page_rect: fitz.Rect,
    pixmap_width: int,
    pixmap_height: int,
) -> fitz.Rect:
    x0, y0, x1, y1 = bbox_pixels
    scale_x = page_rect.width / pixmap_width
    scale_y = page_rect.height / pixmap_height
    return fitz.Rect(x0 * scale_x, y0 * scale_y, (x1 + 1) * scale_x, (y1 + 1) * scale_y)


def _expanded_rect(rect: fitz.Rect, bounds: fitz.Rect, margin: float) -> fitz.Rect:
    return fitz.Rect(
        max(bounds.x0, rect.x0 - margin),
        max(bounds.y0, rect.y0 - margin),
        min(bounds.x1, rect.x1 + margin),
        min(bounds.y1, rect.y1 + margin),
    )


def _standard_page_for_rect(rect: fitz.Rect) -> fitz.Rect:
    aspect = rect.width / max(rect.height, 1.0)
    if aspect >= 1.2:
        return fitz.Rect(0, 0, 1190, 841)
    return fitz.Rect(0, 0, 595, 841)


SHEET_BLOCK_INSERT_OFFSET_X = 420.0


def _entities_for_sheet(
    modelspace: ezdxf.layouts.Modelspace,
    entity_boxes: list[tuple[ezdxf.entities.DXFEntity, BoundingBox2d]],
    sheet_box: BoundingBox2d,
) -> list[ezdxf.entities.DXFEntity]:
    selected: list[ezdxf.entities.DXFEntity] = []
    seen: set[int] = set()

    for entity, entity_box in entity_boxes:
        if not _boxes_intersect(entity_box, sheet_box, margin=2.0):
            continue
        entity_id = id(entity)
        if entity_id in seen:
            continue
        selected.append(entity)
        seen.add(entity_id)

    anchor_x = float(sheet_box.extmin.x) - SHEET_BLOCK_INSERT_OFFSET_X
    anchor_y = float(sheet_box.extmin.y)
    for entity in modelspace:
        if entity.dxftype() != "INSERT":
            continue
        entity_id = id(entity)
        if entity_id in seen:
            continue
        insert = entity.dxf.insert
        if (
            abs(float(insert.x) - anchor_x) <= 1.5
            and abs(float(insert.y) - anchor_y) <= 1.5
        ):
            selected.append(entity)
            seen.add(entity_id)
    return selected


def _entity_render_extents(entity: ezdxf.entities.DXFEntity) -> BoundingBox2d | None:
    try:
        if entity.dxftype() == "INSERT":
            virtual_entities = list(entity.virtual_entities())
            if not virtual_entities:
                return None
            return bbox.extents(virtual_entities)
        if entity.dxftype() == "ACAD_TABLE":
            virtual_entities = list(entity.virtual_entities())
            if not virtual_entities:
                return None
            return bbox.extents(virtual_entities)
        return bbox.extents([entity])
    except Exception:
        return None


def _entity_boxes(modelspace: ezdxf.layouts.Modelspace) -> list[tuple[ezdxf.entities.DXFEntity, BoundingBox2d]]:
    entity_boxes: list[tuple[ezdxf.entities.DXFEntity, BoundingBox2d]] = []
    for entity in modelspace:
        extents = _entity_render_extents(entity)
        if extents is not None and extents.has_data:
            entity_boxes.append((entity, extents))
    return entity_boxes


def _boxes_intersect(a: BoundingBox2d, b: BoundingBox2d, margin: float = 0.0) -> bool:
    return not (
        a.extmax.x < b.extmin.x - margin
        or a.extmin.x > b.extmax.x + margin
        or a.extmax.y < b.extmin.y - margin
        or a.extmin.y > b.extmax.y + margin
    )


def _page_for_layout(
    dxf_layout: ezdxf.layouts.Layout,
    render_box: BoundingBox2d | None = None,
) -> layout.Page:
    extents = render_box if render_box is not None else bbox.extents(dxf_layout)
    width = abs(float(extents.size.x))
    height = abs(float(extents.size.y))
    if not _is_valid_size(width, height):
        return layout.Page(210, 297, units=layout.Units.mm)

    if width >= height:
        if width > 450 or width / max(height, 1.0) > 1.75:
            return layout.Page(420, 297, units=layout.Units.mm)
        return layout.Page(297, 210, units=layout.Units.mm)

    if height > 450 or height / max(width, 1.0) > 1.75:
        return layout.Page(297, 420, units=layout.Units.mm)
    return layout.Page(210, 297, units=layout.Units.mm)


def _modelspace_sheet_boxes(modelspace: ezdxf.layouts.Modelspace) -> list[BoundingBox2d]:
    return _sheet_boxes_for_layout(modelspace, layer_filter={"П"})


def _layout_sheet_boxes(dxf_layout: ezdxf.layouts.Layout) -> list[BoundingBox2d]:
    if dxf_layout.name.lower() == "model":
        return []
    return _sheet_boxes_for_layout(dxf_layout, layer_filter=None)


def _sheet_boxes_for_layout(
    dxf_layout: ezdxf.layouts.Layout,
    layer_filter: set[str] | None,
) -> list[BoundingBox2d]:
    boxes: list[tuple[float, float, float, float]] = []
    for entity in dxf_layout:
        if entity.dxftype() not in {"LWPOLYLINE", "POLYLINE", "INSERT"}:
            continue
        if layer_filter is not None and getattr(entity.dxf, "layer", "") not in layer_filter:
            continue
        try:
            extents = bbox.extents([entity])
        except Exception:
            continue
        min_x = float(extents.extmin.x)
        min_y = float(extents.extmin.y)
        max_x = float(extents.extmax.x)
        max_y = float(extents.extmax.y)
        width = abs(max_x - min_x)
        height = abs(max_y - min_y)
        if _is_sheet_size(width, height):
            boxes.append((min_x, min_y, max_x, max_y))

    unique: dict[tuple[int, int, int, int], tuple[float, float, float, float]] = {}
    for box in boxes:
        key = tuple(round(value, 2) for value in box)
        unique[key] = box

    return [
        BoundingBox2d([Vec2(min_x, min_y), Vec2(max_x, max_y)])
        for min_x, min_y, max_x, max_y in sorted(unique.values(), key=lambda box: (-box[3], box[0]))
    ]


def _is_sheet_size(width: float, height: float) -> bool:
    candidates = ((210.0, 297.0), (297.0, 210.0), (420.0, 297.0), (297.0, 420.0), (594.0, 420.0), (420.0, 594.0))
    return any(abs(width - candidate_width) < 1.5 and abs(height - candidate_height) < 1.5 for candidate_width, candidate_height in candidates)


def _has_finite_extents(dxf_layout: ezdxf.layouts.Layout) -> bool:
    try:
        extents = bbox.extents(dxf_layout)
    except Exception:
        return False
    return _is_valid_size(abs(float(extents.size.x)), abs(float(extents.size.y)))


def _is_valid_size(width: float, height: float) -> bool:
    return all(math.isfinite(value) and value > 1.0 for value in (width, height))
