from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document


@dataclass(frozen=True)
class TuContent:
    text: str
    paragraphs: list[str]
    table_cells: list[str]
    warnings: list[str] | None = None


def read_tu_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".docx":
        return _read_docx(path).text
    if extension == ".pdf":
        return _read_pdf(path).text
    raise ValueError("ТУ должен быть в формате DOCX или PDF.")


def parse_tu(path: Path) -> tuple[dict[str, Any], list[str]]:
    extension = path.suffix.lower()
    if extension == ".docx":
        content = _read_docx(path)
    elif extension == ".pdf":
        content = _read_pdf(path)
    else:
        raise ValueError("ТУ должен быть в формате DOCX или PDF.")

    text = content.text
    warnings: list[str] = list(content.warnings or [])
    if len(_normalize(text)) < 20:
        warnings.append("PDF/ТУ выглядит как скан или почти не содержит извлекаемого текста. Для точного извлечения данных нужен OCR или DOCX-версия ТУ.")

    power_kw = _extract_power_kw(text)
    cable_section = _extract_cable_section(text)
    if not cable_section:
        cable_section = _infer_cable_section_from_power(power_kw)
        if cable_section:
            warnings.append(f"Сечение кабеля не найдено прямой фразой в ТУ. Для MVP принято {cable_section} по максимальной мощности {power_kw} кВт.")

    tp_name = _extract_tp_name(text)
    data = {
        "APPLICANT": _extract_applicant(text, content.paragraphs),
        "ADRESS": _extract_address(text),
        "KADNUMBER": _extract_kadnumber(text),
        "PROJECTNUMBER": _extract_tu_project_number(text, content.table_cells),
        "DATE": _extract_date(text, content.table_cells),
        "POWER_KW": power_kw,
        "VOLTAGE": _extract_voltage(text),
        "FAZE": _extract_phase(text),
        "STROIDOM": _extract_stroidom(text),
        "SECH_KABEL": cable_section,
        "OTKUDASTROIT": _extract_build_from(text, tp_name),
        "TP_NAME": tp_name,
        "requires_komapparat_template": _requires_komapparat_template(text),
        "source_file": path.name,
        "text_length": len(text),
    }

    for key in (
        "APPLICANT",
        "ADRESS",
        "KADNUMBER",
        "DATE",
        "FAZE",
        "STROIDOM",
        "SECH_KABEL",
        "OTKUDASTROIT",
        "TP_NAME",
    ):
        if not data.get(key):
            warnings.append(f"Поле {key} не удалось уверенно извлечь из ТУ.")

    return data, warnings


def build_branch_armature_line(tu_data: dict[str, Any]) -> str:
    tp_name = _normalize(str(tu_data.get("TP_NAME", "") or "")).strip()
    otkuda = _normalize(str(tu_data.get("OTKUDASTROIT", "") or "")).strip()
    if not tp_name:
        return ""

    if "коммутационного аппарата" in otkuda.lower():
        return f"Монтаж ответвительной арматуры от коммутационного аппарата {tp_name}"

    pole_match = re.search(
        r"от\s+опоры\s*№\s*([\w/-]+)\s*ВЛИ?\s*0[,.]4\s*кВ\s*фидера\s*№\s*(\d+)",
        otkuda,
        flags=re.IGNORECASE,
    )
    if pole_match:
        pole, feeder = pole_match.groups()
        return (
            f"Монтаж ответвительной арматуры от сущ. опоры №{pole.strip()} "
            f"ВлИ 0,4 кВ фидера №{feeder} {tp_name}"
        )

    on_pole_match = re.search(r"на\s+опоре\s*№\s*(\d+)", otkuda, flags=re.IGNORECASE)
    if on_pole_match:
        return (
            f"Монтаж ответвительной арматуры на опоре №{on_pole_match.group(1)} "
            f"ВЛ 0,4 кВ фидера №2 {tp_name}"
        )

    return f"Монтаж ответвительной арматуры от коммутационного аппарата {tp_name}"


def _read_docx(path: Path) -> TuContent:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = _normalize(cell.text)
                if value:
                    table_cells.append(value)
    text = "\n".join([*paragraphs, *table_cells])
    return TuContent(text=text, paragraphs=paragraphs, table_cells=table_cells)


def _read_pdf(path: Path) -> TuContent:
    try:
        import fitz  # type: ignore

        document = fitz.open(path)
        text = "\n".join(page.get_text("text") for page in document)
        if len(_normalize(text)) >= 20:
            return TuContent(text=text, paragraphs=[], table_cells=[])

        ocr_text, ocr_warnings = _ocr_pdf_with_tesseract(path, document)
        if ocr_text:
            ocr_warnings.insert(0, "PDF распознан через OCR. Проверьте точность извлечённых данных.")
            return TuContent(text=ocr_text, paragraphs=_ocr_paragraphs(ocr_text), table_cells=[], warnings=ocr_warnings)

        return TuContent(text=text, paragraphs=[], table_cells=[], warnings=ocr_warnings)
    except ModuleNotFoundError:
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            return TuContent(text=text, paragraphs=[], table_cells=[])
        except ModuleNotFoundError as exc:
            raise RuntimeError("Для чтения PDF нужен PyMuPDF или pdfplumber.") from exc


def _ocr_pdf_with_tesseract(path: Path, document: Any) -> tuple[str, list[str]]:
    executable = shutil.which("tesseract")
    if not executable:
        return "", ["PDF выглядит как скан, но tesseract OCR не найден. Загрузите DOCX-версию ТУ или установите tesseract."]

    pages_text: list[str] = []
    warnings: list[str] = []
    with tempfile.TemporaryDirectory(prefix="tu_ocr_") as tmp:
        temp_dir = Path(tmp)
        for index, page in enumerate(document):
            image_path = temp_dir / f"page_{index + 1}.png"
            pixmap = page.get_pixmap(dpi=300, alpha=False)
            pixmap.save(image_path)

            command = [
                executable,
                str(image_path),
                "stdout",
                "-l",
                "rus+eng",
                "--psm",
                "6",
            ]
            try:
                completed = subprocess.run(command, capture_output=True, text=True, timeout=120, check=False)
            except subprocess.TimeoutExpired:
                warnings.append(f"OCR страницы {index + 1} не завершился за 120 секунд.")
                continue

            if completed.returncode != 0:
                message = _normalize(completed.stderr) or f"код {completed.returncode}"
                warnings.append(f"OCR страницы {index + 1} завершился с ошибкой: {message}")
                continue

            pages_text.append(completed.stdout)

    return "\n".join(pages_text), warnings


def _ocr_paragraphs(text: str) -> list[str]:
    return [_normalize(line) for line in text.splitlines() if _normalize(line)]


def _normalize(text: str) -> str:
    text = _fix_ocr_word_breaks(text)
    return re.sub(r"\s+", " ", text).strip()


def _search(pattern: str, text: str, flags: int = re.IGNORECASE | re.DOTALL) -> str:
    match = re.search(pattern, text, flags)
    if not match:
        return ""
    value = next((group for group in match.groups() if group), "")
    return _normalize(value.strip(" .;,\n\t"))


def _extract_applicant(text: str, paragraphs: list[str]) -> str:
    applicant_markers = (
        "фамилия, имя, отчество заявителя",
        "полное наименование организации",
    )
    for index, paragraph in enumerate(paragraphs):
        lower = paragraph.lower()
        if any(marker in lower for marker in applicant_markers):
            previous = _previous_content_paragraph(paragraphs, index)
            if previous:
                return _normalize_applicant_name(previous)

    patterns = [
        r"наименование сетевой организации[^)]*\)\s*\n\s*([^\n(]+?)\s*\n\s*\(полное наименование организации",
        r"\)\s*([^()\n]{5,160}?)\s*\(фамилия,\s*имя,\s*отчество заявителя\)",
        r"\(([^()]{5,120})\),\s*расположенн",
    ]
    for pattern in patterns:
        value = _search(pattern, text)
        if value:
            return _normalize_applicant_name(value)
    return ""


def _normalize_applicant_name(value: str) -> str:
    normalized = abbreviate_applicant_display_terms(_normalize(value))
    lower = normalized.lower()
    organization_markers = (
        "общество",
        "ооо",
        "оао",
        "пао",
        "зао",
        "ип ",
        "администрация",
        "учреждение",
        "компания",
        "«",
        '"',
    )
    if any(marker in lower for marker in organization_markers):
        return normalized

    parts = re.findall(r"[А-ЯЁ][а-яё]+", normalized)
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return normalized


def _extract_address(text: str) -> str:
    item_3 = _numbered_item(text, "3", "4")
    if item_3:
        value = _address_from_item(item_3)
        if value:
            return value

    patterns = [
        r"расположенн\w*\s+по\s+адресу[:\s]+(.+?)(?:;?\s*кадастровый|,?\s*к/н|\.?\s*4\.|\n\s*Основан|\n\s*Общие|\n\s*Исполнен)",
        r"место\s+нахождения.+?[:\s]+(.+?)(?:;?\s*кадастровый|,?\s*к/н|\.?\s*4\.|\n\s*Основан|\n\s*Общие|\n\s*Исполнен)",
    ]
    for pattern in patterns:
        value = _search(pattern, text)
        if value:
            return _trim_address(value)
    return ""


def _extract_kadnumber(text: str) -> str:
    value = _search(r"(\d{2}\s*:\s*\d{2}\s*:\s*\d{6,10}\s*:\s*\d+)", text)
    return re.sub(r"\s+", "", value) if value else ""


def _extract_tu_project_number(text: str, table_cells: list[str]) -> str:
    for value in table_cells:
        project_number = _search(r"\((20\d{2}/\s*\d{3,})\)", value)
        if project_number:
            return project_number
    return _search(r"\b(20\d{2}/\s*\d{3,})\b", text)


def _extract_date(text: str, table_cells: list[str]) -> str:
    for value in table_cells:
        date = _extract_tu_date_from_line(value)
        if date:
            return date

    for line in text.splitlines():
        date = _extract_tu_date_from_line(line)
        if date:
            return date

    for match in re.finditer(r"\b(\d{2}\.\d{2}\.\d{4})\b", text):
        context = text[max(0, match.start() - 80) : match.end() + 80].lower()
        if any(marker in context for marker in ("пп рф", "правил", "утвержден")):
            continue
        return match.group(1)
    return ""


def _extract_power_kw(text: str) -> str:
    item_9 = _numbered_item(text, "9", "10")
    value = _last_kw_value(item_9)
    if value:
        return value

    item_4 = _numbered_item(text, "4", "5")
    value = _last_kw_value(item_4)
    if value:
        return value

    value = _search(r"составляет[:\s]+([\d,.]+)\s*кВт", text)
    if value:
        return value.replace(",", ".")

    value = _search(
        r"максимальн\w*\s+присоединяем\w*\s+мощност\w*\s*[-–—]\s*([\d,.]+)\s*кВт",
        text,
    )
    return value.replace(",", ".") if value else ""


def _last_kw_value(text: str) -> str:
    values = re.findall(r"(\d+(?:[,.]\d+)?)\s*кВт", text or "", re.IGNORECASE)
    return values[-1].replace(",", ".") if values else ""


def _extract_voltage(text: str) -> str:
    item_7 = _numbered_item(text, "7", "8")
    value = _search(r"(\d+(?:[,.]\d+)?)\s*кВ", item_7 or text)
    return value.replace(",", ".") if value else ""


def _extract_phase(text: str) -> str:
    priority_sources = [
        _numbered_item(text, "13.1.6", "13.1.7"),
        _numbered_item(text, "9", "10"),
    ]
    for source in priority_sources:
        value = _phase_from_text(source)
        if value:
            return value
    return _phase_from_text(text)


def _phase_from_text(text: str) -> str:
    lower = text.lower()
    three_phase_markers = (
        r"установить\s+тр[её]хфазн",
        r"тр[её]хфазн\w*\s+прибор",
        r"тр[её]хфазн\w*\s+счетчик",
        r"тр[её]хфазн\w*\s+сч[её]тчик",
        r"тр[её]хфазн",
    )
    single_phase_markers = (
        r"установить\s+однофазн",
        r"однофазн\w*\s+прибор",
        r"однофазн\w*\s+счетчик",
        r"однофазн\w*\s+сч[её]тчик",
        r"однофазн",
    )
    first_three = _first_marker_position(lower, three_phase_markers)
    first_single = _first_marker_position(lower, single_phase_markers)
    if first_three is None and first_single is None:
        return ""
    if first_three is not None and (first_single is None or first_three <= first_single):
        return "трехфазный"
    return "однофазный"


def _first_marker_position(text: str, patterns: tuple[str, ...]) -> int | None:
    positions = [match.start() for pattern in patterns if (match := re.search(pattern, text))]
    return min(positions) if positions else None


def _extract_stroidom(text: str) -> str:
    item_1 = _numbered_item(text, "1", "2")
    value = _search(
        r"вводное\s+устройство\s*\(ВУ\)\s*(.+?)(?:\.|;|$)",
        item_1,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if value:
        return value

    item_3 = _numbered_item(text, "3", "4")
    value = _search(
        r"заявителя[:\s]+([^:.;,\n]+?)(?::|;|,|\.)",
        item_3,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if value:
        return value

    value = _search(r"устройств\s+заявителя:\s*([^,\n(]+)", text, flags=re.IGNORECASE)
    return _normalize(value) if value else ""


def _extract_cable_section(text: str) -> str:
    lower = text.lower().replace("х", "x")
    if re.search(r"3\s*x\s*70\s*\+\s*1\s*x\s*70", lower):
        return "3x70+1x70"
    if re.search(r"3\s*x\s*50\s*\+\s*1\s*x\s*50", lower):
        return "3x50+1x50"
    if re.search(r"от\s*50[^.\n;]{0,40}до\s*(70|100)", lower):
        return "3x70+1x70"
    if re.search(r"до\s*50", lower):
        return "3x50+1x50"
    return ""


def _extract_build_from(text: str, tp_name: str = "") -> str:
    for number, next_number in (("13.2.1", "13.2.2"), ("13.1.1", "13.1.2")):
        item = _numbered_item(text, number, next_number)
        if _starts_from_commutator(item):
            return _commutator_build_phrase(tp_name)
        value = _build_fragment_from_item(item)
        if value:
            return value

    value = _search(
        r"(от\s+опоры\s*№\s*.+?)(?:\s+построить|\s+смонтировать|\s+ориентировочной|$)",
        text,
    )
    if value:
        return _fix_ocr_word_breaks(value).strip(" .;,")
    return ""


def _extract_tp_name(text: str) -> str:
    return _search(r"(ТП\s*№\s*[\wА-Яа-яЁё/\- ]+?\s*кВА)", text)


def _starts_from_commutator(value: str) -> bool:
    normalized = _normalize(value).lower().replace("ё", "е")
    return bool(re.search(r"\bот\s+коммутационного\s+аппарата\b", normalized))


def _commutator_build_phrase(tp_name: str) -> str:
    tp_name = _normalize(tp_name)
    if tp_name:
        return f"от коммутационного аппарата {tp_name} до границы участка Заявителя"
    return "от коммутационного аппарата до границы участка Заявителя"


def _requires_komapparat_template(text: str) -> bool:
    item_13_1_3 = _numbered_item(text, "13.1.3", "13.1.4")
    if item_13_1_3:
        return _has_additional_komapparat_requirement(item_13_1_3)

    match = re.search(
        r"(?ms)(?:^|\n)\s*13\.1\.3\.\s*(.+?)(?=\n\s*13\.1\.4\.|\Z)",
        text,
    )
    return bool(match and _has_additional_komapparat_requirement(match.group(1)))


def _has_additional_komapparat_requirement(value: str) -> bool:
    normalized = _normalize(value).lower().replace("ё", "е")
    if "коммутацион" not in normalized or "аппарат" not in normalized:
        return False
    return any(marker in normalized for marker in ("смонтировать", "установить", "монтаж"))


def _trim_address(value: str) -> str:
    markers = (
        "Липецкая область",
        "г. ",
        "город ",
        "деревня ",
        "село ",
        "поселок ",
        "посёлок ",
    )
    for marker in markers:
        index = value.find(marker)
        if index >= 0:
            return _cut_address_tail(value[index:])
    return _cut_address_tail(value)


def _previous_content_paragraph(paragraphs: list[str], index: int) -> str:
    for candidate in reversed(paragraphs[:index]):
        lower = candidate.lower()
        if candidate.startswith("(") and candidate.endswith(")"):
            continue
        if "наименование сетевой организации" in lower:
            continue
        if "россети" in lower or "липецкэнерго" in lower:
            continue
        if candidate:
            return candidate
    return ""


def _numbered_item(text: str, number: str, next_number: str) -> str:
    pattern = (
        rf"(?ms)(?:^|\n)\s*{re.escape(number)}\.\s*(.+?)"
        rf"(?=\n\s*(?:{re.escape(next_number)}|\d+(?:\.\d+)*)\.|\Z)"
    )
    return _search(pattern, text, flags=0)


def _address_from_item(item: str) -> str:
    lipetsk_index = item.find("Липецкая область")
    if lipetsk_index >= 0:
        return _cut_address_tail(item[lipetsk_index:])

    for marker in ("г. ", "город ", "деревня ", "село ", "поселок ", "посёлок "):
        index = item.find(marker)
        if index >= 0:
            return _cut_address_tail(item[index:])
    return ""


def _cut_address_tail(value: str) -> str:
    result = value
    for marker in ("кадастровый номер", "к/н", "кадастровый"):
        index = result.lower().find(marker)
        if index >= 0:
            result = result[:index]
    return _normalize_address_terms(result.strip(" .;,"))


def _extract_tu_date_from_line(value: str) -> str:
    match = re.search(
        r"«\s*_*\s*(\d{1,2})\s*_*\s*»\s*_*\s*(\d{1,2})\s*_*\s*(\d{4})\s*[гr]",
        value,
        re.IGNORECASE,
    )
    if not match:
        return ""

    day, month, year = match.groups()
    return f"{int(day):02d}.{int(month):02d}.{year}"


def _build_fragment_from_item(item: str) -> str:
    if not item:
        return ""
    value = _search(r"(от\s+.+?)(?:\s+смонтировать|\s+построить|\s+ориентировочной|$)", item)
    return _fix_ocr_word_breaks(value).strip(" .;,")


def _fix_ocr_word_breaks(value: str) -> str:
    return re.sub(r"(?<=[А-Яа-яЁё])-\s+(?=[а-яё])", "", value)


_GARDEN_PARTNERSHIP_PATTERNS = (
    r"(?:[А-ЯЁ][а-яё-]+\s+)?садоводческое\s+(?:некоммерческое\s+партнерство|потребительское\s+общество|некоммерческое\s+товарищество)",
    r"(?:[А-ЯЁ][а-яё-]+\s+)?садовое\s+(?:некоммерческое\s+партнерство|потребительское\s+общество|некоммерческое\s+товарищество)",
    r"\bсадоводческое\s+некоммерческое\s+партнерство\b",
    r"\bсадовое\s+некоммерческое\s+партнерство\b",
)

_GARAGE_COOPERATIVE_PATTERNS = (
    (r"\bгаражный\s+потребительский\s+кооператив\s+автолюбителей\b", "ГПКА"),
    (r"\bгаражный\s+потребительский\s+кооператив\b", "ГПК"),
    (r"\bгаражный\s+кооператив\b", "ГК"),
)

_APPLICANT_ORGANIZATION_PATTERNS = (
  (r"\bобщество\s+с\s+ограниченной\s+ответственностью\b", "ООО"),
  (r"\bиндивидуальный\s+предприниматель\b", "ИП"),
  (r"\bакционерное\s+общество\b", "АО"),
  (r"\bпубличное\s+акционерное\s+общество\b", "ПАО"),
)


def abbreviate_applicant_display_terms(value: str) -> str:
    result = value
    for pattern, replacement in _APPLICANT_ORGANIZATION_PATTERNS:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", result).strip()


def abbreviate_garden_partnership_terms(value: str) -> str:
    return abbreviate_address_display_terms(value)


def abbreviate_address_display_terms(value: str) -> str:
    result = value
    for pattern in _GARDEN_PARTNERSHIP_PATTERNS:
        result = re.sub(pattern, "СНП", result, flags=re.IGNORECASE)
    for pattern, replacement in _GARAGE_COOPERATIVE_PATTERNS:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    result = re.sub(r"\bучасток\b", "уч.", result, flags=re.IGNORECASE)
    result = _drop_redundant_urban_district(result)
    return re.sub(r"\s+", " ", result).strip()


def _drop_redundant_urban_district(value: str) -> str:
    return re.sub(
        r",\s*городской\s+округ\s+(г\.?\s*[А-ЯЁа-яё-]+)\s*,(?=.*\1)",
        ", ",
        value,
        flags=re.IGNORECASE,
    )


def _normalize_address_terms(value: str) -> str:
    has_city_marker = bool(re.search(r"\b(?:г\.|город)\s*[А-ЯЁA-Z]", value, flags=re.IGNORECASE))
    result = _normalize(value)
    result = re.sub(
        r"\b([А-ЯЁа-яё-]+?ский)\s+муниципальный\s+район\b",
        r"\1 район",
        result,
        flags=re.IGNORECASE,
    )
    result = re.sub(
        r"\b([А-ЯЁа-яё-]+?ского)\s+муниципального\s+района\b",
        r"\1 района",
        result,
        flags=re.IGNORECASE,
    )
    result = re.sub(
        r"\b([А-ЯЁ][а-яё]+ский)\s+муниципальный\s+округ\b",
        r"\1 округ",
        result,
        flags=re.IGNORECASE,
    )
    result = re.sub(
        r"\b([А-ЯЁа-яё-]+?ского)\s+муниципального\s+округа\b",
        r"\1 округа",
        result,
        flags=re.IGNORECASE,
    )
    result = re.sub(r"\bсельское\s+поселение\b", "с.п.", result, flags=re.IGNORECASE)
    replacements = (
        (r"\bгород\s+", "г."),
        (r"\bсело\s+", "с."),
        (r"\bдеревня\s+", "д."),
        (r"\bпос[её]лок\s+", "п."),
        (r"\bулица\s+", "ул."),
        (r"\bдом\s+", "д."),
    )
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    result = re.sub(r"\b(г|с|д|п|ул)\.\s+", r"\1.", result, flags=re.IGNORECASE)
    result = re.sub(r"\bс\.п\.\s+", "с.п. ", result, flags=re.IGNORECASE)
    result = re.sub(r"\bс\.п\.(?=[А-ЯЁ])", "с.п. ", result)
    result = re.sub(r"\bземельный участок\b", "з.у.", result, flags=re.IGNORECASE)
    result = re.sub(r"\s*,\s*", ", ", result)
    if has_city_marker:
        result = _drop_district_before_city(result)
    result = abbreviate_address_display_terms(result)
    return result.strip(" .;,")


def _drop_district_before_city(value: str) -> str:
    return re.sub(
        r",\s*[^,]*?\bрайон(?:а)?\b\s*,\s*(?=г\.)",
        ", ",
        value,
        flags=re.IGNORECASE,
    )


def _infer_cable_section_from_power(power_kw: str) -> str:
    try:
        power = float(str(power_kw).replace(",", "."))
    except ValueError:
        return ""

    if power > 50:
        return "3x70+1x70"
    if power > 0:
        return "3x50+1x50"
    return ""
